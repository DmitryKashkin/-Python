from docx import Document
from os import mkdir


def clear_path(path_temp: str, bad_symbols: list):
    # Функция заменяет недопустимые символы в имени файла (пути) на пробел
    for i in bad_symbols:
        path_temp = path_temp.replace(i, ' ')
    return path_temp


def create_path(row):
    # Функция генерирует из табличной строки путь (имя файла)
    bad_symbols = ['\n', ":", "?"]
    path = row.cells[0].text + '-' + row.cells[1].text + ' ' + row.cells[3].text
    path = clear_path(path, bad_symbols)
    return path


def length_normalization(list_of_path: list):
    # Функция ограничивает длину имени файла (пути) до 90 символов
    for i in range(len(list_of_path)):
        if len(list_of_path[i]) > 90:
            list_of_path[i] = list_of_path[i][:90]
        while list_of_path[-1] == ' ' or list_of_path[-1] == '.':  # TODO не работает: не удаляет висячий пробел!!!
            list_of_path.pop()


def write_path(list_of_path: list):
    doc = Document()
    prefix = 'c:/temp/'
    length_normalization(list_of_path)
    for path in list_of_path:
        filename = path + '.docx'
        path = prefix + path
        filename = path + "/" + filename
        print(path)
        print(filename)

        mkdir(path)
        doc.save(filename)


def main():
    doc = Document('calendar_list.docx')
    # последовательность всех таблиц документа
    all_tables = doc.tables
    table = all_tables[0]
    list_of_articles = table.columns[3]
    list_of_path = []
    for row in table.rows:
        temp_path = create_path(row)
        list_of_path.append(temp_path)
    list_of_path = list_of_path[1:]
    write_path(list_of_path)


def create_set_of_all_symbols(list_of_articles):
    all_symbols = set()  # в это множество добавим все символы из названий, чтобы вычислить левые символы
    for cell in list_of_articles.cells:
        all_symbols.update(set(cell.text))
    all_symbols = list(all_symbols)
    all_symbols.sort()
    print(all_symbols)


if __name__ == '__main__':
    main()

